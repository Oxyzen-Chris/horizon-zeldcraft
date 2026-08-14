# -*- coding: utf-8 -*-
"""Génère Horizon_ZeldCraft_Financement_Communication.docx à la racine du repo.
Usage : python docs/marketing/build_docx.py
Nécessite : pip install python-docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

GOLD = RGBColor(0xB8, 0x86, 0x0B)
PURPLE = RGBColor(0x4B, 0x2E, 0x83)
DARK = RGBColor(0x1E, 0x1B, 0x2E)

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT_PATH = os.path.join(ROOT, "Horizon_ZeldCraft_Financement_Communication.docx")

doc = Document()

# ---- styles de base ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for i, size in enumerate([28, 20, 15, 12], start=1):
    try:
        h = doc.styles[f"Heading {i}"]
        h.font.size = Pt(size)
        h.font.color.rgb = PURPLE if i <= 2 else DARK
        h.font.bold = True
    except KeyError:
        pass


def add_title_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🗡️ HORIZON ZELDCRAFT")
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = GOLD

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Le Tamagotchi Web3 qui grandit avec toi")
    run2.font.size = Pt(18)
    run2.italic = True
    run2.font.color.rgb = PURPLE

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(
        "Dossier de communication & de financement participatif"
    )
    run3.font.size = Pt(14)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Instagram : @horizon.zeldcraft  •  github.com/Oxyzen-Chris/horizon-zeldcraft")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_page_break()


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def add_bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 4"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


# =====================================================================
add_title_page()

# ---------------------------------------------------------------------
add_heading("1. En une phrase", level=1)
add_para(
    "Horizon ZeldCraft est un compagnon virtuel — Synk, un petit héros pixel-art — que l'on "
    "nourrit et fait grandir chaque jour, chaque semaine, chaque mois et chaque année, comme un "
    "Tamagotchi, mais dont la croissance est enregistrée pour de vrai sur la blockchain Ethereum : "
    "chaque repas est une petite preuve numérique infalsifiable, à toi pour toujours."
)

add_heading("2. L'histoire, sans jargon technique", level=1)
add_para(
    "Synk est un jeune aventurier, un peu à la manière du héros de Zelda, habillé façon Minecraft. "
    "Il commence son épopée les mains vides dans la forêt de Zephyria, et grandit au fil de tes "
    "attentions et de tes exploits : il devient guerrier, puis mage, puis dresseur de dragons. "
    "Son monde, Zeldcraftia, mélange l'ambiance de Minecraft Dungeons, la douceur de Zelda : Breath "
    "of the Wild / Tears of the Kingdom, et l'épique de World of Warcraft."
)
add_para(
    "Le grand méchant de l'histoire s'appelle Zorghon : il a enlevé la princesse PocaPoka et son "
    "fidèle lutin El Pipo. Toute l'aventure — 400 quêtes réparties en 40 chapitres — consiste à "
    "les retrouver et à vaincre Zorghon, en réunissant 5 fragments d'un sceau magique légendaire. "
    "En chemin : des PNJ à qui parler, des îles à explorer en bateau ou en engin volant, des "
    "montagnes enneigées, des saisons qui changent avec le calendrier réel, et même des nuits de "
    "pleine lune qui révèlent des quêtes secrètes."
)
add_para(
    "Premier compagnon apprivoisable : un Dragon d'Or, le plus noble des dragons, débloqué à force "
    "de bons soins. D'autres créatures rejoindront Synk au fil des saisons de contenu."
)

add_heading("3. Pourquoi c'est différent d'un jeu mobile classique", level=1)
add_bullets([
    "Ton compagnon t'appartient réellement : sa croissance est inscrite sur la blockchain "
    "Ethereum, pas seulement sur le serveur d'une entreprise qui peut fermer du jour au lendemain.",
    "Tu commences gratuitement sur le réseau de test Sepolia (ETH gratuits via un « robinet »/"
    "faucet), sans dépenser un centime, pour découvrir le jeu.",
    "Le jeu est jouable sans même posséder de portefeuille crypto (mode « Jouer sans "
    "portefeuille » ou « Accès démo ») — l'univers Zelda/Minecraft/WoW reste accessible à tous.",
    "Connexion possible avec MetaMask, Rainbow, Coinbase, WalletConnect, Ledger, ou simplement "
    "avec un compte e-mail/mot de passe.",
    "Disponible sur navigateur (Vercel) et sur mobile via Expo Go — un seul compte, deux façons "
    "d'y jouer.",
    "Multijoueur en tête : classements de joueurs, rencontres, économie partagée — comme "
    "Minecraft Dungeons, pensé solo au départ puis ouvert à plusieurs.",
])

add_heading("4. Un jeu construit différemment : l'histoire derrière le code", level=1)
add_para(
    "Deux façons de raconter comment Horizon ZeldCraft a vu le jour — à choisir selon le public "
    "visé (à tester en A/B sur les réseaux) :"
)
add_para("Angle A — « Le 1er jeu assisté à 100% par une IA »", bold=True)
add_para(
    "Horizon ZeldCraft a été conçu avec GitHub Copilot comme partenaire de développement de bout "
    "en bout : smart contract, interface de jeu, back-office d'administration, notifications e-mail, "
    "tests automatisés — le tout piloté par le dialogue entre un développeur et une IA. Un pari sur "
    "la nouvelle façon de fabriquer des jeux vidéo."
)
add_para("Angle B — « Un jeu complet livré en 3 semaines, 3-4h par jour »", bold=True)
add_para(
    "Horizon ZeldCraft, de l'idée au jeu jouable en ligne (contrat Ethereum, interface, back-office, "
    "e-mails, multilingue), a été construit en environ 3 semaines à raison de 3-4 heures par jour : "
    "rédaction de consignes précises à une IA, tests, déploiements, correctifs de bugs, recherche "
    "d'inspiration, fidélité à l'esprit Donjons & Dragons."
)
add_para(
    "Recommandation : privilégier l'angle B (« construit en 3 semaines ») pour le grand public — "
    "c'est concret, mesurable, et valorise le savoir-faire du développeur autant que l'outil. "
    "Réserver l'angle A (« 100% IA ») aux publics tech/Web3 (Twitter/X, communautés dev), où "
    "l'expérimentation « IA-first » suscite davantage la curiosité.",
    italic=True,
)

add_heading("5. Ce qui arrive ensuite (feuille de route)", level=1)
add_bullets([
    "🌱 Printemps / ☀️ Été / 🍂 Automne / ❄️ Hiver — de nouvelles quêtes et rencontres à chaque "
    "saison, activables d'un simple interrupteur en coulisses.",
    "Nouveaux familiers (dragons chromatiques et métalliques, elfes des forêts) et nouvelles îles "
    "à explorer.",
    "Passage progressif du réseau de test Sepolia vers le réseau principal Ethereum, pour jouer "
    "avec de vrais ETH en toute confiance (après audit de sécurité du contrat).",
    "Marketplace communautaire pour échanger skins et objets entre joueurs.",
])

add_heading("6. Pourquoi on a besoin d'un coup de main", level=1)
add_para(
    "Horizon ZeldCraft tourne aujourd'hui grâce à des services gratuits (hébergement web, base de "
    "données, envoi d'e-mails). C'est parfait pour les débuts, mais si la communauté grandit, ces "
    "services deviennent payants — et le passage sur le réseau Ethereum principal (avec de vrais "
    "ETH) demande une vérification de sécurité du contrat avant de manipuler de l'argent réel."
)
add_table(
    ["Poste de dépense", "Aujourd'hui", "Une fois le jeu populaire"],
    [
        ["Hébergement web (Vercel)", "Gratuit (Hobby)", "≈ 20 $/mois (Pro)"],
        ["Base de données (Firebase)", "Gratuit (Spark)", "Variable selon trafic (Blaze, à l'usage)"],
        ["Envoi d'e-mails (Resend)", "Gratuit (bac à sable)", "≈ 20 $/mois pour un vrai domaine + volume"],
        ["Nom de domaine", "—", "≈ 10-15 €/an"],
        ["Accès blockchain (RPC Alchemy/Infura)", "Gratuit (petit volume)", "≈ 50-100 $/mois si succès"],
        ["Audit de sécurité du contrat (Mainnet)", "—", "Poste le plus important, souvent 1 500-5 000 $"],
        ["Frais de gas (déploiement Mainnet)", "—", "Variable selon le prix du gas ETH"],
        ["Temps de développement & support joueurs", "Bénévole", "À valoriser si le projet grandit"],
    ],
)
add_para(
    "L'objectif du sponsoring n'est pas de « faire fortune », mais de mutualiser ces coûts : "
    "chaque petite contribution, cumulée à celle de tous les joueurs et sponsors, permet de garder "
    "le jeu en ligne, sûr, et de continuer à l'enrichir.",
    italic=True,
)

add_heading("7. Comment nous soutenir", level=1)
add_para(
    "Trois canaux gratuits à mettre en place, complémentaires, sans passer par une levée de fonds "
    "classique :"
)
add_table(
    ["Plateforme", "Type de soutien", "Contreparties possibles"],
    [
        [
            "Ko-fi (recommandé pour démarrer)",
            "Don ponctuel ou abonnement mensuel",
            "Skin exclusif, badge « Fondateur » in-game, mention au générique",
        ],
        [
            "GitHub Sponsors",
            "Don ponctuel ou mensuel, public technique",
            "Accès anticipé aux nouvelles fonctionnalités, mention dans le repo",
        ],
        [
            "Open Collective",
            "Don avec suivi transparent des dépenses",
            "Visibilité complète sur l'usage des fonds (infra, audit, dev)",
        ],
        [
            "Gitcoin Grants (piste à moyen terme)",
            "Financement communautaire Web3 avec matching de fonds",
            "Visibilité auprès de la communauté Ethereum",
        ],
    ],
)
add_para(
    "Paliers de soutien suggérés (Ko-fi) : 3€ « Aventurier » (badge in-game), 10€/mois "
    "« Chevalier » (skin exclusif + accès prioritaire aux nouveautés), 25€/mois « Maître Dresseur "
    "de Dragons » (mention au générique + vote consultatif sur le prochain familier de saison)."
)

add_heading("8. Plan de communication Instagram (@horizon.zeldcraft)", level=1)
add_para(
    "Voir le calendrier détaillé et prêt à l'emploi dans "
    "docs/marketing/instagram-content-plan.md (8 premiers posts, légendes et hashtags inclus), "
    "ainsi que les visuels réels du jeu dans docs/marketing/screenshots/ et la vidéo teaser "
    "docs/marketing/teaser-fr.mp4."
)
add_bullets([
    "Rythme conseillé : 2-3 publications par semaine au lancement, en alternant capture de jeu, "
    "lore (présentation de Synk, Zorghon, les dragons), et coulisses du développement.",
    "Formats : visuels carrés issus du jeu, courtes vidéos de gameplay, citations de la lore.",
    "Cross-promotion suggérée : sous-communautés Reddit (r/WGaming, r/ethdev, r/IndieGaming), "
    "Discord communautaire dédié, X/Twitter pour le public Web3.",
])
add_para(
    "⚠️ Note de transparence : les cinématiques 3D « cinéma » avec dragons animés crachant du feu "
    "pour dessiner le logo, ou traveling de caméra dans un château, nécessitent un vrai pipeline "
    "d'animation 3D (par ex. Blender, gratuit mais demandant un travail de modélisation/animation "
    "manuel important) — hors de portée d'une automatisation sans budget. La vidéo teaser fournie "
    "ici est un montage honnête à partir de vraies captures d'écran du jeu, avec effets de "
    "mouvement de caméra (Ken Burns) et incrustations de texte, qui peut déjà servir de démo "
    "d'accueil dans le jeu en attendant un budget dédié à une vraie cinématique.",
)

add_heading("9. Contact", level=1)
add_para("Instagram : @horizon.zeldcraft")
add_para("GitHub : github.com/Oxyzen-Chris/horizon-zeldcraft")
add_para("Site : horizon-zeldcraft.vercel.app")

doc.save(OUT_PATH)
print(f"Document généré : {OUT_PATH}")
